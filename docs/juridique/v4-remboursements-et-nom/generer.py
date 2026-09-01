# -*- coding: utf-8 -*-
"""Met la note du 1er septembre au format Word, dans la mise en forme des documents
déjà relus — mêmes polices, mêmes couleurs, mêmes marges.

La source reste le Markdown : c'est lui qu'on modifie, jamais le Word. Deux
fichiers qui disent la même chose finissent toujours par diverger, et c'est
celui qu'on a sous les yeux qui a tort.

    python generer.py
"""

import io
import os
import re

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ACCENT = RGBColor(0xA3, 0x4D, 0x1F)
ENCRE = RGBColor(0x1A, 0x16, 0x14)
GRIS = RGBColor(0x5C, 0x56, 0x50)

ICI = os.path.dirname(os.path.abspath(__file__))
SOURCE = os.path.join(ICI, 'NOTE-AU-JURISTE-v4.md')
CIBLE = os.path.join(ICI, 'NOTE-AU-JURISTE-v4.docx')


def style_de_base(doc):
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(11)
    normal.font.color.rgb = ENCRE
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(7)

    for section in doc.sections:
        section.top_margin = Inches(0.87)
        section.bottom_margin = Inches(0.87)
        section.left_margin = Inches(0.98)
        section.right_margin = Inches(0.98)


def ecrire_riche(paragraphe, texte):
    """Rend `**gras**` et `` `code` `` sans dépendre d'une bibliothèque de plus.

    Le Markdown de ces notes n'emploie que ces deux marques : un analyseur
    complet coûterait une dépendance pour deux cas.
    """
    for morceau in re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', texte):
        if not morceau:
            continue
        if morceau.startswith('**') and morceau.endswith('**'):
            passage = paragraphe.add_run(morceau[2:-2])
            passage.bold = True
        elif morceau.startswith('`') and morceau.endswith('`'):
            passage = paragraphe.add_run(morceau[1:-1])
            passage.font.name = 'Consolas'
            passage.font.size = Pt(9.5)
            passage.font.color.rgb = GRIS
        else:
            paragraphe.add_run(morceau.replace('_', ''))


def ajouter_tableau(doc, lignes):
    entetes = [c.strip() for c in lignes[0].strip('|').split('|')]
    corps = [
        [c.strip() for c in ligne.strip('|').split('|')]
        for ligne in lignes[2:]
    ]

    tableau = doc.add_table(rows=1, cols=len(entetes))
    tableau.style = 'Table Grid'

    for cellule, titre in zip(tableau.rows[0].cells, entetes):
        cellule.text = ''
        p = cellule.paragraphs[0]
        passage = p.add_run(titre)
        passage.bold = True
        passage.font.color.rgb = ACCENT
        passage.font.size = Pt(10)

    for ligne in corps:
        cellules = tableau.add_row().cells
        for cellule, valeur in zip(cellules, ligne):
            cellule.text = ''
            p = cellule.paragraphs[0]
            ecrire_riche(p, valeur)
            for passage in p.runs:
                passage.font.size = Pt(10)

    doc.add_paragraph()


def rejoindre_les_replis(lignes):
    """Recolle les lignes repliées à quatre-vingt-dix colonnes.

    Le Markdown est écrit en colonnes étroites pour rester lisible ; sans cette
    passe, chaque repli devient un paragraphe Word distinct et une puce se
    coupe en deux au milieu d'une phrase.

    Une ligne est un repli si la précédente porte du texte et qu'elle-même ne
    commence par aucune marque — titre, puce, tuyau, citation, trait.
    """
    marques = ('#', '-', '*', '|', '>')
    sorties = []

    for ligne in lignes:
        nu = ligne.strip()
        debute_un_bloc = (
            nu == ''
            or nu.startswith(marques)
            or re.match(r'^\d+\. ', nu) is not None
        )

        if sorties and sorties[-1].strip() and not debute_un_bloc:
            sorties[-1] = sorties[-1].rstrip() + ' ' + nu
        else:
            sorties.append(ligne)

    return sorties


def convertir():
    lignes = rejoindre_les_replis(
        io.open(SOURCE, encoding='utf-8').read().split('\n')
    )

    doc = Document()
    style_de_base(doc)

    entete = doc.add_paragraph()
    marque = entete.add_run('PasseTemps')
    marque.bold = True
    marque.font.size = Pt(13)
    marque.font.color.rgb = ACCENT

    sous = doc.add_paragraph()
    detail = sous.add_run('anciennement Ma Villa · hébergements et artisanat — Sénégal')
    detail.font.size = Pt(9)
    detail.font.color.rgb = GRIS

    i = 0
    while i < len(lignes):
        ligne = lignes[i].rstrip()

        # Tableau : une ligne de tuyaux suivie d'un séparateur.
        if ligne.startswith('|') and i + 1 < len(lignes) and set(lignes[i + 1].replace('|', '').strip()) <= set('-: '):
            bloc = []
            while i < len(lignes) and lignes[i].startswith('|'):
                bloc.append(lignes[i])
                i += 1
            ajouter_tableau(doc, bloc)
            continue

        if ligne.startswith('# '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            passage = p.add_run(ligne[2:])
            passage.bold = True
            passage.font.size = Pt(18)
            passage.font.color.rgb = ENCRE

        elif ligne.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(16)
            passage = p.add_run(ligne[3:])
            passage.bold = True
            passage.font.size = Pt(13)
            passage.font.color.rgb = ACCENT

        elif ligne.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(11)
            passage = p.add_run(ligne[4:])
            passage.bold = True
            passage.font.size = Pt(11.5)
            passage.font.color.rgb = ENCRE

        elif ligne.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            ecrire_riche(p, ligne[2:])
            for passage in p.runs:
                passage.font.color.rgb = GRIS

        elif re.match(r'^\d+\. ', ligne):
            p = doc.add_paragraph(style='List Number')
            ecrire_riche(p, re.sub(r'^\d+\. ', '', ligne))

        elif ligne.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            ecrire_riche(p, ligne[2:])

        elif ligne.startswith('---'):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            trait = p.add_run('• • •')
            trait.font.color.rgb = GRIS

        elif ligne.strip():
            p = doc.add_paragraph()
            ecrire_riche(p, ligne)

        i += 1

    doc.save(CIBLE)
    return CIBLE


if __name__ == '__main__':
    print('Écrit :', convertir())
