# -*- coding: utf-8 -*-
"""Produit les CGU version 2 au format Word, dans la mise en forme des documents
déjà relus par le juriste — mêmes polices, mêmes couleurs, mêmes marges.

Le texte reprend mot pour mot la version validée le 12 août 2026, sauf là où le
logiciel a changé. Chaque section touchée porte une mention, afin que la
relecture porte sur ce qui bouge et non sur l'ensemble.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ACCENT = RGBColor(0xA3, 0x4D, 0x1F)
ENCRE = RGBColor(0x1A, 0x16, 0x14)
GRIS = RGBColor(0x5C, 0x56, 0x50)

TITRE = "Conditions générales d'utilisation"
VERSION = "Version 2 — 18 août 2026 · introduction de l'encaissement en ligne"

CHAPEAU = (
    "Cette version remplace celle du 10 août. Elle n'a pas valeur d'avis juridique. "
    "Trois sections ont été réécrites et deux sont nouvelles, parce que la plateforme "
    "encaisse désormais les paiements : elles portent la mention « modifié », "
    "« réécrit » ou « nouveau ». Les autres sont inchangées et déjà relues. Les "
    "passages entre crochets appellent une décision."
)

INTRO = (
    "Les présentes conditions régissent l'utilisation de la plateforme Ma Villa. "
    "En créant un compte, vous les acceptez sans réserve."
)

SECTIONS = [
    ("1. Objet et rôle de la plateforme", "modifié", [
        "Ma Villa, exploitée par [raison sociale à préciser], est une plateforme de mise en "
        "relation entre des propriétaires de villas et de logements de vacances situés au "
        "Sénégal et des clients souhaitant les louer.",

        "Ma Villa agit en qualité d'intermédiaire technique. Elle n'est ni propriétaire, ni "
        "gestionnaire, ni loueur des biens présentés. Le contrat de location est conclu "
        "directement entre le client et le propriétaire.",

        "Ma Villa encaisse toutefois le prix des réservations au nom et pour le compte du "
        "propriétaire, qui lui en donne mandat par son inscription. Cet encaissement ne fait "
        "pas de Ma Villa une partie au contrat de location.",

        "Ma Villa ne saurait être tenue responsable de l'état réel du bien, de sa conformité "
        "à l'annonce, ni de l'exécution du séjour.",
    ]),

    ("2. Comptes utilisateurs", None, [
        "L'inscription est gratuite. Trois rôles existent : client, propriétaire, "
        "administrateur.",

        "L'utilisateur s'engage à fournir des informations exactes et à maintenir la "
        "confidentialité de ses identifiants. Toute activité réalisée depuis son compte lui "
        "est imputable.",

        "Ma Villa peut suspendre ou supprimer un compte en cas de manquement aux présentes "
        "conditions, notamment en cas d'annonce mensongère, d'avis frauduleux ou de "
        "comportement abusif.",
    ]),

    ("3. Publication d'une annonce", None, [
        "Le propriétaire garantit qu'il détient les droits nécessaires pour louer le bien "
        "publié et que les informations diffusées (description, photographies, tarifs, "
        "capacité, localisation) sont exactes et à jour.",

        "Toute annonce est soumise à validation par l'équipe Ma Villa avant sa mise en ligne. "
        "Une annonce peut être rejetée ou retirée sans préavis si elle est incomplète, "
        "trompeuse ou contraire à la loi.",

        "Le propriétaire est seul responsable du respect de ses obligations légales et "
        "fiscales liées à l'activité de location.",
    ]),

    ("4. Réservations", "modifié", [
        "Une demande de réservation précise le logement, la formule tarifaire, les dates et "
        "le nombre de personnes.",

        "Une réservation réglée en ligne est confirmée dès l'encaissement du paiement, sans "
        "intervention du propriétaire. Le propriétaire s'engage à honorer toute réservation "
        "ainsi confirmée.",

        "Une réservation non réglée en ligne n'engage définitivement les parties qu'après "
        "confirmation par le propriétaire, qui s'engage à répondre dans un délai raisonnable. "
        "Une demande sans réponse ne vaut pas acceptation.",

        "Le nombre de personnes déclaré ne peut excéder la capacité du logement. Tout "
        "dépassement autorise le propriétaire à refuser l'accès.",
    ]),

    ("5. Prix, paiement et commission", "réécrit", [
        "Les tarifs sont fixés librement par les propriétaires et affichés en francs CFA "
        "(FCFA), toutes taxes comprises lorsque celles-ci sont applicables.",

        "Le paiement s'effectue en ligne, par Wave ou Orange Money, au moyen des services du "
        "prestataire PayDunya. Ma Villa n'a accès à aucune donnée bancaire ni à aucun code "
        "secret : ceux-ci sont saisis dans l'application du moyen de paiement choisi.",

        "Le montant payé par le client est exactement celui affiché sur l'annonce. Aucun "
        "frais de service ne s'y ajoute.",

        "Sur ce montant, Ma Villa retient une commission au titre de la mise en relation et "
        "de la gestion de l'encaissement : 20 % pour une réservation d'un montant supérieur "
        "ou égal à 50 000 FCFA, 10 % pour une réservation d'un montant inférieur. La "
        "commission est arrondie à l'unité inférieure, l'écart profitant au propriétaire.",

        "Le propriétaire perçoit le solde, soit 80 % ou 90 % du montant payé selon le cas.",

        "Le prestataire de paiement n'accepte aucune transaction inférieure à 200 FCFA. En "
        "deçà, le règlement s'effectue directement entre le client et le propriétaire.",

        "Ma Villa peut modifier ses taux de commission. Toute modification est portée à la "
        "connaissance des propriétaires [délai de préavis à préciser] avant son entrée en "
        "vigueur et ne s'applique qu'aux réservations postérieures.",
    ]),

    ("6. Reversement au propriétaire", "nouveau", [
        "Les sommes encaissées pour le compte du propriétaire, déduction faite de la "
        "commission, lui sont reversées [délai et périodicité à préciser] après la fin du "
        "séjour.",

        "Le reversement s'effectue sur le compte Wave ou Orange Money déclaré par le "
        "propriétaire. Celui-ci est responsable de l'exactitude de ces coordonnées ; Ma Villa "
        "ne répond pas d'un versement adressé à des coordonnées erronées qu'il aurait "
        "lui-même renseignées.",

        "Ma Villa peut suspendre un reversement en cas de litige déclaré, de soupçon de "
        "fraude, ou d'annulation ouvrant droit à remboursement, le temps d'instruire le "
        "dossier.",

        "Le propriétaire dispose du détail de ses réservations et des montants correspondants "
        "dans son espace personnel.",
    ]),

    ("7. Annulation et remboursement", "nouveau", [
        "Les conditions d'annulation figurent dans la politique d'annulation, qui fait partie "
        "intégrante des présentes conditions.",

        "Lorsque la réservation a été réglée en ligne, le remboursement éventuel est effectué "
        "par Ma Villa sur le moyen de paiement d'origine, dans le délai indiqué par cette "
        "politique.",

        "[À trancher : en cas de remboursement intégral, la commission est-elle restituée au "
        "client ou conservée au titre du service rendu ? Le choix doit être énoncé ici sans "
        "ambiguïté, car il détermine le montant effectivement remboursé.]",
    ]),

    ("8. Avis", None, [
        "Seul un client ayant effectué un séjour confirmé et terminé dans la villa concernée "
        "peut déposer un avis. Un avis par client et par villa.",

        "Les avis engagent leur auteur. Ma Villa se réserve le droit de retirer tout avis "
        "injurieux, diffamatoire, hors sujet ou manifestement frauduleux.",
    ]),

    ("9. Responsabilité", "modifié", [
        "Ma Villa met en œuvre les moyens raisonnables pour assurer la disponibilité et la "
        "sécurité de la plateforme, sans garantie d'absence d'interruption ou d'erreur.",

        "La responsabilité de Ma Villa ne peut être engagée en cas de litige entre un client "
        "et un propriétaire, de dommage survenu pendant un séjour, ou d'inexécution imputable "
        "à l'une des parties.",

        "Le service de paiement est fourni par un prestataire tiers. Ma Villa ne répond pas "
        "d'une indisponibilité, d'un retard ou d'un refus imputable à ce prestataire ou à "
        "l'opérateur de paiement mobile choisi par le client.",
    ]),

    ("10. Données personnelles", None, [
        "Le traitement des données personnelles est décrit dans la politique de "
        "confidentialité, qui fait partie intégrante des présentes conditions.",
    ]),

    ("11. Propriété intellectuelle", None, [
        "La marque, le nom de domaine, la charte graphique et les développements de la "
        "plateforme sont la propriété exclusive de son exploitant.",

        "En publiant des photographies, le propriétaire concède à Ma Villa une licence "
        "gratuite et non exclusive d'utilisation à des fins de promotion de son annonce et "
        "de la plateforme.",
    ]),

    ("12. Modification et droit applicable", None, [
        "Ma Villa peut modifier les présentes conditions à tout moment. La version applicable "
        "est celle en vigueur à la date d'utilisation du service.",

        "Les présentes conditions sont soumises au droit sénégalais. À défaut de règlement "
        "amiable, tout litige relève de la compétence des juridictions de Dakar.",
    ]),
]

RETOUR = (
    "Merci de retourner ce fichier annoté ou corrigé, afin que la version validée soit "
    "intégrée telle quelle au site."
)


def texte(doc, contenu, taille=None, couleur=None, gras=False, italique=False,
          avant=0, apres=6, alignement=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(avant)
    p.paragraph_format.space_after = Pt(apres)
    if alignement is not None:
        p.alignment = alignement
    r = p.add_run(contenu)
    if taille:
        r.font.size = Pt(taille)
    if couleur:
        r.font.color.rgb = couleur
    r.bold = gras
    r.italic = italique
    return p


def construire(chemin):
    doc = Document()

    s = doc.sections[0]
    s.left_margin = s.right_margin = Inches(0.98)
    s.top_margin = s.bottom_margin = Inches(0.87)

    normal = doc.styles['Normal'].font
    normal.name = 'Calibri'
    normal.size = Pt(11)

    # En-tête
    texte(doc, 'MA VILLA', taille=13, couleur=ACCENT, gras=True, apres=0)
    texte(doc, 'Plateforme de location de villas au Sénégal', taille=9, couleur=GRIS, apres=14)

    t = doc.add_paragraph(style='Title')
    rt = t.add_run(TITRE)
    rt.font.color.rgb = ENCRE

    texte(doc, VERSION, taille=9, couleur=GRIS, apres=14)

    texte(doc, "DOCUMENT SOUMIS À RELECTURE JURIDIQUE — NON PUBLIÉ EN L'ÉTAT",
          taille=9.5, gras=True, apres=4)
    texte(doc, CHAPEAU, taille=9.5, couleur=GRIS, apres=16)

    texte(doc, INTRO, apres=12)

    for titre, mention, paragraphes in SECTIONS:
        h = doc.add_paragraph(style='Heading 1')
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(4)
        rh = h.add_run(titre)
        rh.font.size = Pt(12.5)
        rh.font.color.rgb = ACCENT

        if mention:
            rm = h.add_run('   — ' + mention)
            rm.font.size = Pt(9)
            rm.font.color.rgb = GRIS
            rm.italic = True

        for para in paragraphes:
            texte(doc, para, apres=6)

    texte(doc, '— Fin du document —', couleur=GRIS, avant=18, apres=4,
          alignement=WD_ALIGN_PARAGRAPH.CENTER)
    texte(doc, RETOUR, taille=9, couleur=GRIS,
          alignement=WD_ALIGN_PARAGRAPH.CENTER)

    doc.save(chemin)
    return len(SECTIONS)


if __name__ == '__main__':
    import os
    cible = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                         '1-Conditions-generales-utilisation-v2.docx')
    n = construire(cible)
    print('Écrit : %s (%d sections)' % (cible, n))
